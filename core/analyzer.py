import os
import re
from typing import Dict, List, Any, Optional, Tuple
from docx import Document
import logging
from pathlib import Path

from .utils import setup_logging, load_yaml_config, clean_text, is_docx_file
from .retrieval import DocumentRetriever

logger = setup_logging()

class DocumentAnalyzer:
    def __init__(self, rules_path: str = "rules", config_path: str = "config/settings.yml"):
        self.rules_path = rules_path
        self.config = load_yaml_config(config_path)
        self.retriever = DocumentRetriever()
        
        # Load checklists and red-flag rules
        self.checklists = self._load_checklists()
        self.redflag_rules = self._load_redflag_rules()
        
        # Process detection patterns
        self.process_patterns = {
            "Company Incorporation": [
                r"incorporation",
                r"articles of association",
                r"memorandum of association",
                r"register of members",
                r"register of directors",
                r"ubo declaration",
                r"name reservation"
            ],
            "Employment": [
                r"employment contract",
                r"employee handbook",
                r"terms of employment",
                r"er 2024",
                r"employment regulations"
            ],
            "Post Registration": [
                r"articles amendment",
                r"shareholder resolution",
                r"board resolution",
                r"change of directors",
                r"change of registered office"
            ],
            "Annual Filings": [
                r"annual accounts",
                r"annual return",
                r"annual filing",
                r"financial statements"
            ]
        }
    
    def _load_checklists(self) -> Dict[str, Any]:
        """Load all checklist YAML files"""
        checklists = {}
        checklists_dir = Path(self.rules_path) / "checklists"
        
        if checklists_dir.exists():
            for yaml_file in checklists_dir.glob("*.yml"):
                if yaml_file.name != "redflags":  # Skip redflags directory
                    try:
                        checklist = load_yaml_config(str(yaml_file))
                        process_name = checklist.get("process", yaml_file.stem)
                        checklists[process_name] = checklist
                    except Exception as e:
                        logger.error(f"Error loading checklist {yaml_file}: {e}")
        
        return checklists
    
    def _load_redflag_rules(self) -> Dict[str, Any]:
        """Load red-flag rules"""
        redflags_dir = Path(self.rules_path) / "checklists" / "redflags"
        rules = {}
        
        if redflags_dir.exists():
            for yaml_file in redflags_dir.glob("*.yml"):
                try:
                    rule_set = load_yaml_config(str(yaml_file))
                    scope = rule_set.get("scope", yaml_file.stem)
                    rules[scope] = rule_set
                except Exception as e:
                    logger.error(f"Error loading redflag rules {yaml_file}: {e}")
        
        return rules
    
    def parse_docx(self, file_path: str) -> Dict[str, Any]:
        """Parse a .docx file and extract text and structure"""
        if not is_docx_file(file_path):
            raise ValueError(f"File {file_path} is not a .docx file")
        
        try:
            doc = Document(file_path)
            
            # Extract text from paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append({
                        'text': para.text.strip(),
                        'style': para.style.name if para.style else 'Normal'
                    })
            
            # Extract text from tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    table_data.append(row_data)
                tables.append(table_data)
            
            # Combine all text
            full_text = "\n".join([p['text'] for p in paragraphs])
            
            return {
                'file_path': file_path,
                'filename': Path(file_path).name,
                'paragraphs': paragraphs,
                'tables': tables,
                'full_text': full_text,
                'word_count': len(full_text.split()),
                'paragraph_count': len(paragraphs),
                'table_count': len(tables)
            }
            
        except Exception as e:
            logger.error(f"Error parsing docx file {file_path}: {e}")
            raise
    
    def detect_process(self, documents: List[Dict[str, Any]]) -> str:
        """Detect the process type from uploaded documents"""
        combined_text = " ".join([doc.get('full_text', '') for doc in documents])
        combined_text = combined_text.lower()
        
        # Score each process type
        process_scores = {}
        
        for process_name, patterns in self.process_patterns.items():
            score = 0
            for pattern in patterns:
                matches = re.findall(pattern, combined_text, re.IGNORECASE)
                score += len(matches)
            process_scores[process_name] = score
        
        # Return the process with highest score
        if process_scores:
            best_process = max(process_scores.items(), key=lambda x: x[1])
            if best_process[1] > 0:
                return best_process[0]
        
        # Default fallback
        return "General Review"
    
    def detect_entity_type(self, documents: List[Dict[str, Any]]) -> str:
        """Detect entity type from documents"""
        combined_text = " ".join([doc.get('full_text', '') for doc in documents])
        combined_text = combined_text.lower()
        
        # Entity type detection patterns
        entity_patterns = {
            "Private Company Limited by Shares (Non-Financial)": [
                r"private company limited by shares",
                r"limited by shares",
                r"share capital",
                r"shares issued"
            ],
            "Private Company Limited by Guarantee (Non-Financial)": [
                r"limited by guarantee",
                r"guarantee company",
                r"no share capital"
            ],
            "Branch (Non-Financial)": [
                r"branch office",
                r"branch registration",
                r"foreign company"
            ]
        }
        
        for entity_type, patterns in entity_patterns.items():
            for pattern in patterns:
                if re.search(pattern, combined_text, re.IGNORECASE):
                    return entity_type
        
        return "Private Company Limited by Shares (Non-Financial)"  # Default
    
    def check_redflags(self, documents: List[Dict[str, Any]], 
                      process: str, entity_type: str) -> List[Dict[str, Any]]:
        """Check for red flags in documents"""
        redflags = []
        
        # Get applicable red-flag rules
        applicable_rules = []
        for scope, rule_set in self.redflag_rules.items():
            if scope == "General Corporate Docs" or scope.lower() in process.lower():
                applicable_rules.extend(rule_set.get("rules", {}).items())
        
        for rule_name, rule_config in applicable_rules:
            rule_redflags = self._apply_redflag_rule(
                rule_name, rule_config, documents, process, entity_type
            )
            redflags.extend(rule_redflags)
        
        return redflags
    
    def _apply_redflag_rule(self, rule_name: str, rule_config: Dict[str, Any],
                           documents: List[Dict[str, Any]], process: str, 
                           entity_type: str) -> List[Dict[str, Any]]:
        """Apply a specific red-flag rule to documents"""
        redflags = []
        
        # Check if rule applies
        applies_if = rule_config.get("applies_if", "always")
        if applies_if != "always":
            if applies_if == "process == 'Company Incorporation'" and process != "Company Incorporation":
                return redflags
            # Add more conditions as needed
        
        # Check entity type conditions
        trigger_if = rule_config.get("trigger_if", [])
        if trigger_if:
            entity_condition_met = False
            for condition in trigger_if:
                if "entity_type ==" in condition:
                    required_entity = condition.split("==")[1].strip().strip("'\"")
                    if entity_type == required_entity:
                        entity_condition_met = True
                        break
            if not entity_condition_met:
                return redflags
        
        # Apply rule based on type
        rule_kind = rule_config.get("kind", "")
        applies_to_docs = rule_config.get("applies_to_docs", ["All"])
        
        for doc in documents:
            doc_name = doc.get('filename', 'Unknown')
            
            # Check if rule applies to this document
            if "All" not in applies_to_docs and doc_name not in applies_to_docs:
                continue
            
            doc_text = doc.get('full_text', '')
            
            if rule_kind == "pattern_presence":
                redflag = self._check_pattern_presence_rule(
                    rule_name, rule_config, doc_name, doc_text
                )
                if redflag:
                    redflags.append(redflag)
            
            elif rule_kind == "structural_check":
                redflag = self._check_structural_rule(
                    rule_name, rule_config, doc_name, doc_text
                )
                if redflag:
                    redflags.append(redflag)
            
            elif rule_kind == "semantic_check":
                redflag = self._check_semantic_rule(
                    rule_name, rule_config, doc_name, doc_text
                )
                if redflag:
                    redflags.append(redflag)
            
            elif rule_kind == "heuristic":
                redflag = self._check_heuristic_rule(
                    rule_name, rule_config, doc_name, doc_text
                )
                if redflag:
                    redflags.append(redflag)
        
        return redflags
    
    def _check_pattern_presence_rule(self, rule_name: str, rule_config: Dict[str, Any],
                                   doc_name: str, doc_text: str) -> Optional[Dict[str, Any]]:
        """Check pattern presence rule"""
        patterns_any = rule_config.get("patterns_any", [])
        require_phrase = rule_config.get("require_phrase", "")
        
        # Check if any forbidden patterns are present
        forbidden_found = False
        for pattern in patterns_any:
            if re.search(pattern, doc_text, re.IGNORECASE):
                forbidden_found = True
                break
        
        if forbidden_found:
            # Check if required phrase is missing
            if require_phrase and not re.search(require_phrase, doc_text, re.IGNORECASE):
                return {
                    "rule": rule_name,
                    "document": doc_name,
                    "issue": rule_config.get("message", "Pattern presence issue detected"),
                    "severity": rule_config.get("severity", "Medium"),
                    "citations": rule_config.get("citations", [])
                }
        
        return None
    
    def _check_structural_rule(self, rule_name: str, rule_config: Dict[str, Any],
                             doc_name: str, doc_text: str) -> Optional[Dict[str, Any]]:
        """Check structural rule"""
        checks = rule_config.get("checks", [])
        
        # Simple structural checks
        missing_checks = []
        
        if "has_signatory_name" in checks:
            if not re.search(r"(signed|signature|signed by|executed by).*\b[A-Z][a-z]+ [A-Z][a-z]+\b", doc_text, re.IGNORECASE):
                missing_checks.append("signatory name")
        
        if "has_capacity" in checks:
            if not re.search(r"(director|officer|authorized|capacity)", doc_text, re.IGNORECASE):
                missing_checks.append("capacity")
        
        if "has_signature_or_e-sign" in checks:
            if not re.search(r"(signature|signed|electronic signature|e-sign)", doc_text, re.IGNORECASE):
                missing_checks.append("signature")
        
        if "has_date" in checks:
            if not re.search(r"\b\d{1,2}[/-]\d{1,2}[/-]\d{4}\b|\b\d{4}[/-]\d{1,2}[/-]\d{1,2}\b", doc_text):
                missing_checks.append("date")
        
        if missing_checks:
            return {
                "rule": rule_name,
                "document": doc_name,
                "issue": f"Missing structural elements: {', '.join(missing_checks)}",
                "severity": rule_config.get("severity", "Medium"),
                "citations": rule_config.get("citations", [])
            }
        
        return None
    
    def _check_semantic_rule(self, rule_name: str, rule_config: Dict[str, Any],
                           doc_name: str, doc_text: str) -> Optional[Dict[str, Any]]:
        """Check semantic rule"""
        forbidden_phrases = rule_config.get("forbidden_phrases", [])
        
        for phrase in forbidden_phrases:
            if re.search(phrase, doc_text, re.IGNORECASE):
                return {
                    "rule": rule_name,
                    "document": doc_name,
                    "issue": rule_config.get("message", f"Forbidden phrase found: {phrase}"),
                    "severity": rule_config.get("severity", "High"),
                    "citations": rule_config.get("citations", [])
                }
        
        return None
    
    def _check_heuristic_rule(self, rule_name: str, rule_config: Dict[str, Any],
                            doc_name: str, doc_text: str) -> Optional[Dict[str, Any]]:
        """Check heuristic rule"""
        indicators = rule_config.get("indicators_any", [])
        
        for indicator in indicators:
            if "Template fields left blank" in indicator:
                if re.search(r"\[\[.*?\]\]|_{3,}", doc_text):
                    return {
                        "rule": rule_name,
                        "document": doc_name,
                        "issue": rule_config.get("message", "Template placeholders found"),
                        "severity": rule_config.get("severity", "Low"),
                        "citations": rule_config.get("citations", [])
                    }
            
            elif "Lorem ipsum" in indicator:
                if "lorem ipsum" in doc_text.lower():
                    return {
                        "rule": rule_name,
                        "document": doc_name,
                        "issue": rule_config.get("message", "Lorem ipsum text found"),
                        "severity": rule_config.get("severity", "Low"),
                        "citations": rule_config.get("citations", [])
                    }
        
        return None
    
    def analyze_documents(self, file_paths: List[str]) -> Dict[str, Any]:
        """Complete document analysis pipeline"""
        logger.info(f"Analyzing {len(file_paths)} documents")
        
        # Parse all documents
        documents = []
        for file_path in file_paths:
            try:
                doc_data = self.parse_docx(file_path)
                documents.append(doc_data)
            except Exception as e:
                logger.error(f"Error parsing {file_path}: {e}")
                continue
        
        if not documents:
            raise ValueError("No valid documents to analyze")
        
        # Detect process and entity type
        process = self.detect_process(documents)
        entity_type = self.detect_entity_type(documents)
        
        # Check red flags
        redflags = self.check_redflags(documents, process, entity_type)
        
        return {
            "documents": documents,
            "process": process,
            "entity_type": entity_type,
            "redflags": redflags,
            "document_count": len(documents)
        }
