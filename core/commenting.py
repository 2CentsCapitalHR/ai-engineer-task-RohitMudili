import os
import re
from typing import Dict, List, Any, Optional, Tuple
from docx import Document
from docx.shared import RGBColor, Inches
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.shared import OxmlElement, qn
import logging
from pathlib import Path
from datetime import datetime

from .utils import setup_logging, load_yaml_config, format_citation

logger = setup_logging()

class DocumentCommenter:
    def __init__(self, config_path: str = "config/settings.yml"):
        self.config = load_yaml_config(config_path)
        self.commenting_mode = self.config.get('commenting', {}).get('mode', 'inline')
        
        # Color schemes for different severity levels
        self.severity_colors = {
            "High": RGBColor(255, 0, 0),      # Red
            "Medium": RGBColor(255, 165, 0),  # Orange
            "Low": RGBColor(255, 255, 0)      # Yellow
        }
    
    def add_comments_to_document(self, file_path: str, issues: List[Dict[str, Any]], 
                                output_path: Optional[str] = None) -> str:
        """Add comments and highlighting to a document based on issues found"""
        try:
            # Load the document
            doc = Document(file_path)
            
            if self.commenting_mode == "aspose":
                return self._add_aspose_comments(doc, issues, output_path)
            else:
                return self._add_inline_comments(doc, issues, output_path)
                
        except Exception as e:
            logger.error(f"Error adding comments to document {file_path}: {e}")
            raise
    
    def _add_inline_comments(self, doc: Document, issues: List[Dict[str, Any]], 
                           output_path: Optional[str] = None) -> str:
        """Add inline comments and highlighting to document"""
        # Group issues by document section for better organization
        issues_by_section = self._group_issues_by_section(issues)
        
        # Add review summary table at the beginning
        self._add_review_summary_table(doc, issues)
        
        # Process each paragraph and add comments where issues are found
        for paragraph in doc.paragraphs:
            self._process_paragraph_for_issues(paragraph, issues)
        
        # Add comments to tables if any
        for table in doc.tables:
            self._process_table_for_issues(table, issues)
        
        # Save the document
        if not output_path:
            input_path = Path(doc._path) if hasattr(doc, '_path') else Path("temp.docx")
            output_path = str(input_path.parent / f"{input_path.stem}_reviewed{input_path.suffix}")
        
        doc.save(output_path)
        logger.info(f"Document with inline comments saved to: {output_path}")
        
        return output_path
    
    def _add_aspose_comments(self, doc: Document, issues: List[Dict[str, Any]], 
                           output_path: Optional[str] = None) -> str:
        """Add comments using Aspose (if available)"""
        # This would use Aspose.Words for Python if available
        # For now, fall back to inline comments
        logger.warning("Aspose mode requested but not implemented. Falling back to inline comments.")
        return self._add_inline_comments(doc, issues, output_path)
    
    def _group_issues_by_section(self, issues: List[Dict[str, Any]]) -> Dict[str, List[Dict[str, Any]]]:
        """Group issues by document section"""
        grouped = {}
        
        for issue in issues:
            section = issue.get('section', 'General')
            if section not in grouped:
                grouped[section] = []
            grouped[section].append(issue)
        
        return grouped
    
    def _add_review_summary_table(self, doc: Document, issues: List[Dict[str, Any]]):
        """Add a review summary table at the beginning of the document"""
        # Add a title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run("Document Review Summary")
        title_run.bold = True
        title_run.font.size = 16
        
        # Add timestamp
        timestamp_para = doc.add_paragraph()
        timestamp_para.add_run(f"Review Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        # Create summary table
        if issues:
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            
            # Header row
            header_cells = table.rows[0].cells
            header_cells[0].text = "Issue"
            header_cells[1].text = "Section"
            header_cells[2].text = "Severity"
            header_cells[3].text = "Suggestion"
            header_cells[4].text = "Citation"
            
            # Make header bold
            for cell in header_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            
            # Add issue rows
            for issue in issues:
                row_cells = table.add_row().cells
                row_cells[0].text = issue.get('issue', '')
                row_cells[1].text = issue.get('section', '')
                row_cells[2].text = issue.get('severity', '')
                row_cells[3].text = issue.get('suggestion', '')
                
                # Add citation
                citations = issue.get('citations', [])
                if citations:
                    row_cells[4].text = citations[0]  # Show first citation
                else:
                    row_cells[4].text = "No citation"
        
        # Add spacing
        doc.add_paragraph()
    
    def _process_paragraph_for_issues(self, paragraph, issues: List[Dict[str, Any]]):
        """Process a paragraph and add comments for relevant issues"""
        paragraph_text = paragraph.text.lower()
        
        for issue in issues:
            # Check if this issue relates to this paragraph
            if self._issue_matches_paragraph(issue, paragraph_text):
                self._highlight_text_in_paragraph(paragraph, issue)
                self._add_inline_comment(paragraph, issue)
    
    def _process_table_for_issues(self, table, issues: List[Dict[str, Any]]):
        """Process a table and add comments for relevant issues"""
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.lower()
                
                for issue in issues:
                    if self._issue_matches_paragraph(issue, cell_text):
                        self._highlight_text_in_cell(cell, issue)
                        self._add_inline_comment_to_cell(cell, issue)
    
    def _issue_matches_paragraph(self, issue: Dict[str, Any], paragraph_text: str) -> bool:
        """Check if an issue matches the content of a paragraph"""
        issue_text = issue.get('issue', '').lower()
        section = issue.get('section', '').lower()
        
        # Check if issue text appears in paragraph
        if issue_text and any(word in paragraph_text for word in issue_text.split()):
            return True
        
        # Check if section reference appears in paragraph
        if section and section in paragraph_text:
            return True
        
        # Check for specific patterns based on issue type
        if "jurisdiction" in issue_text and ("court" in paragraph_text or "law" in paragraph_text):
            return True
        
        if "signature" in issue_text and ("signed" in paragraph_text or "signature" in paragraph_text):
            return True
        
        if "date" in issue_text and re.search(r'\d{1,2}[/-]\d{1,2}[/-]\d{4}', paragraph_text):
            return True
        
        return False
    
    def _highlight_text_in_paragraph(self, paragraph, issue: Dict[str, Any]):
        """Highlight relevant text in a paragraph"""
        severity = issue.get('severity', 'Medium')
        color = self.severity_colors.get(severity, RGBColor(255, 165, 0))
        
        # Find and highlight text that matches the issue
        issue_text = issue.get('issue', '')
        paragraph_text = paragraph.text
        
        # Simple highlighting - could be enhanced with more sophisticated text matching
        if issue_text.lower() in paragraph_text.lower():
            # Clear existing runs and recreate with highlighting
            paragraph.clear()
            
            # Split text and add highlighting
            parts = paragraph_text.split(issue_text)
            for i, part in enumerate(parts):
                if part:
                    run = paragraph.add_run(part)
                
                if i < len(parts) - 1:  # Not the last part
                    highlighted_run = paragraph.add_run(issue_text)
                    highlighted_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    highlighted_run.font.color.rgb = color
    
    def _highlight_text_in_cell(self, cell, issue: Dict[str, Any]):
        """Highlight relevant text in a table cell"""
        severity = issue.get('severity', 'Medium')
        color = self.severity_colors.get(severity, RGBColor(255, 165, 0))
        
        # Similar highlighting logic for table cells
        for paragraph in cell.paragraphs:
            self._highlight_text_in_paragraph(paragraph, issue)
    
    def _add_inline_comment(self, paragraph, issue: Dict[str, Any]):
        """Add an inline comment to a paragraph"""
        comment_text = self._format_comment_text(issue)
        
        # Add comment as a new paragraph with indentation
        comment_para = paragraph._parent.add_paragraph()
        comment_para.style = paragraph.style
        
        # Add comment marker
        comment_run = comment_para.add_run("[Comment: ")
        comment_run.font.color.rgb = RGBColor(128, 128, 128)  # Gray
        comment_run.font.italic = True
        
        # Add comment content
        comment_content = comment_para.add_run(comment_text)
        comment_content.font.color.rgb = RGBColor(0, 0, 255)  # Blue
        comment_content.font.italic = True
        
        # Close comment marker
        comment_end = comment_para.add_run("]")
        comment_end.font.color.rgb = RGBColor(128, 128, 128)  # Gray
        comment_end.font.italic = True
    
    def _add_inline_comment_to_cell(self, cell, issue: Dict[str, Any]):
        """Add an inline comment to a table cell"""
        comment_text = self._format_comment_text(issue)
        
        # Add comment to the cell
        comment_para = cell.add_paragraph()
        comment_run = comment_para.add_run(f"[Comment: {comment_text}]")
        comment_run.font.color.rgb = RGBColor(0, 0, 255)  # Blue
        comment_run.font.italic = True
        comment_run.font.size = 8  # Smaller font for cell comments
    
    def _format_comment_text(self, issue: Dict[str, Any]) -> str:
        """Format the comment text for an issue"""
        parts = []
        
        # Add issue description
        parts.append(issue.get('issue', ''))
        
        # Add suggestion if available
        suggestion = issue.get('suggestion', '')
        if suggestion:
            parts.append(f" Suggestion: {suggestion}")
        
        # Add citation if available
        citations = issue.get('citations', [])
        if citations:
            citation = format_citation(citations[0])
            parts.append(f" Reference: {citation}")
        
        return "".join(parts)
    
    def create_clean_document(self, file_path: str, issues: List[Dict[str, Any]], 
                            output_path: Optional[str] = None) -> str:
        """Create a clean version of the document with comments removed"""
        try:
            doc = Document(file_path)
            
            # Remove any existing comments or highlights
            self._remove_existing_comments(doc)
            
            # Save clean version
            if not output_path:
                input_path = Path(file_path)
                output_path = str(input_path.parent / f"{input_path.stem}_clean{input_path.suffix}")
            
            doc.save(output_path)
            logger.info(f"Clean document saved to: {output_path}")
            
            return output_path
            
        except Exception as e:
            logger.error(f"Error creating clean document: {e}")
            raise
    
    def _remove_existing_comments(self, doc: Document):
        """Remove existing comments and highlights from document"""
        # This is a simplified version - in practice, you'd need more sophisticated
        # logic to identify and remove comments based on the commenting system used
        
        # Remove review summary tables (simple approach)
        paragraphs_to_remove = []
        for paragraph in doc.paragraphs:
            if "Document Review Summary" in paragraph.text:
                paragraphs_to_remove.append(paragraph)
        
        for paragraph in paragraphs_to_remove:
            p = paragraph._element
            p.getparent().remove(p)
    
    def generate_comment_report(self, issues: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Generate a summary report of all comments/issues"""
        report = {
            "total_issues": len(issues),
            "issues_by_severity": {},
            "issues_by_section": {},
            "citations_used": set(),
            "summary": ""
        }
        
        # Count issues by severity
        for issue in issues:
            severity = issue.get('severity', 'Unknown')
            report["issues_by_severity"][severity] = report["issues_by_severity"].get(severity, 0) + 1
            
            section = issue.get('section', 'General')
            report["issues_by_section"][section] = report["issues_by_section"].get(section, 0) + 1
            
            # Collect citations
            citations = issue.get('citations', [])
            report["citations_used"].update(citations)
        
        # Convert set to list for JSON serialization
        report["citations_used"] = list(report["citations_used"])
        
        # Generate summary
        high_issues = report["issues_by_severity"].get("High", 0)
        medium_issues = report["issues_by_severity"].get("Medium", 0)
        low_issues = report["issues_by_severity"].get("Low", 0)
        
        report["summary"] = f"Found {len(issues)} issues: {high_issues} high, {medium_issues} medium, {low_issues} low severity."
        
        return report
